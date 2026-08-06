from docx import Document as DocxDocument
import pandas as pd
import json

import uuid
from pathlib import Path
from typing import List
import pymupdf

from src.core.schema import Document
from src.utils.logger import get_logger

logger = get_logger(__name__)

class DocumentLoader:
    def __init__(self, file_path: str | Path):
        self.file_path = Path(file_path)

    def load(self) -> List[Document]:
        if not self.file_path.exists():
            raise FileNotFoundError(f"File not found: {self.file_path}")

        loaders = {
            ".pdf": self._load_pdf,
            ".txt": self._load_text,
            ".md": self._load_text,
            ".docx": self._load_docx,
            ".csv": self._load_csv,
            ".json": self._load_json,
        }

        loader = loaders.get(self.file_path.suffix.lower())

        if loader is None:
            raise ValueError(
                f"Unsupported file type: {self.file_path.suffix}"
            )
        return loader()

    def _load_pdf(self) -> List[Document]:
        documents = []
        try:
            reader = pymupdf.open(str(self.file_path))
            for i, page in enumerate(reader):
                text = page.get_text()
                text = " ".join(text.split())
                if text:
                    doc = Document(
                        id=str(uuid.uuid4()),
                        content=text,
                        metadata={
                            "source": str(self.file_path.name),
                            "page": i + 1,
                            "type": "pdf"
                        }
                    )
                    documents.append(doc)
        except Exception as e:
            logger.error(f"Error loading PDF {self.file_path.name}: {e}")
        return documents

    def _load_text(self) -> List[Document]:
        try:
            content = self.file_path.read_text(encoding="utf-8")
            doc = Document(
                id=str(uuid.uuid4()),
                content=content,
                metadata={
                    "source": str(self.file_path.name),
                    "type": self.file_path.suffix.lower()[1:]
                }
            )
            return [doc]
        except Exception as e:
            logger.error(f"Error loading text file {self.file_path}: {e}")
            return []

    def _load_docx(self) -> List[Document]:
        try:
            doc = DocxDocument(self.file_path)

            text = "\n".join(
                para.text.strip()
                for para in doc.paragraphs
                if para.text.strip()
            )

            return [
                Document(
                    id=str(uuid.uuid4()),
                    content=text,
                    metadata={
                        "source": self.file_path.name,
                        "type": "docx",
                    },
                )
            ]

        except Exception as e:
            logger.error(f"Error loading DOCX {self.file_path}: {e}")
            return []

    def _load_csv(self) -> List[Document]:
        try:
            df = pd.read_csv(self.file_path)

            text = df.to_string(index=False)

            return [
                Document(
                    id=str(uuid.uuid4()),
                    content=text,
                    metadata={
                        "source": self.file_path.name,
                        "type": "csv",
                        "rows": len(df),
                    },
                )
            ]

        except Exception as e:
            logger.error(f"Error loading CSV {self.file_path}: {e}")
            return []

    def _load_json(self) -> List[Document]:
        try:
            with open(self.file_path, encoding="utf-8") as f:
                obj = json.load(f)

            text = json.dumps(obj, indent=2)

            return [
                Document(
                    id=str(uuid.uuid4()),
                    content=text,
                    metadata={
                        "source": self.file_path.name,
                        "type": "json",
                    },
                )
            ]

        except Exception as e:
            logger.error(f"Error loading JSON {self.file_path}: {e}")
            return []